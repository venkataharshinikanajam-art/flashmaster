"""
Build FLASHMASTER_REPORT.docx from FLASHMASTER_REPORT.md.

Custom parser tuned to the structure of this specific report.
Handles: headings, bold/italic/inline-code, bullets, numbered lists,
fenced code blocks, tables, images, horizontal rules, the centered
cover-page div, and skips Mermaid blocks (Word can't render them).
"""

from __future__ import annotations
import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r"C:/Users/harsh/OneDrive/Desktop/SRM/Full-stack/FLASHCARDS")
SRC = ROOT / "FLASHMASTER_REPORT.md"
OUT = ROOT / "FLASHMASTER_REPORT.docx"


# ----------------------- helpers ---------------------------------------------

INLINE_RE = re.compile(
    r"(\*\*(?P<bold>.+?)\*\*|`(?P<code>[^`]+?)`|\*(?P<ital>[^*]+?)\*|\[(?P<ltext>[^\]]+?)\]\((?P<lurl>[^)]+?)\))"
)


def add_inline(paragraph, text: str, base_bold: bool = False) -> None:
    """Render a text string with inline markdown (bold, italic, code, links)."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            run.bold = base_bold
        if m.group("bold") is not None:
            run = paragraph.add_run(m.group("bold"))
            run.bold = True
        elif m.group("code") is not None:
            run = paragraph.add_run(m.group("code"))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif m.group("ital") is not None:
            run = paragraph.add_run(m.group("ital"))
            run.italic = True
        elif m.group("ltext") is not None:
            run = paragraph.add_run(m.group("ltext"))
            run.bold = base_bold
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.underline = True
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = base_bold


def set_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_horizontal_rule(doc) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pbdr.append(bottom)
    p_pr.append(pbdr)


def add_code_block(doc, code: str, language: str = "") -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    p_pr.append(shd)
    pbdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "DDDDDD")
        pbdr.append(b)
    p_pr.append(pbdr)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


# ----------------------- parser ----------------------------------------------


def parse_and_write(md_text: str, doc: Document) -> None:
    lines = md_text.splitlines()
    i = 0
    in_cover = False
    cover_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        raw = line
        stripped = line.strip()

        # ---- centered cover-page div ------------------------------------
        if stripped.startswith('<div align="center">'):
            in_cover = True
            cover_lines = []
            i += 1
            continue
        if stripped.startswith("</div>") and in_cover:
            in_cover = False
            for cl in cover_lines:
                cs = cl.strip()
                if not cs:
                    doc.add_paragraph()
                    continue
                if cs == "<br>":
                    doc.add_paragraph()
                    continue
                if cs.startswith("# "):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(cs[2:].strip())
                    run.bold = True
                    run.font.size = Pt(28)
                elif cs.startswith("## "):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(cs[3:].strip())
                    run.bold = True
                    run.font.size = Pt(16)
                else:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_inline(p, cs)
            doc.add_page_break()
            i += 1
            continue
        if in_cover:
            cover_lines.append(raw)
            i += 1
            continue

        # ---- horizontal rule --------------------------------------------
        if stripped == "---":
            add_horizontal_rule(doc)
            i += 1
            continue

        # ---- fenced code block ------------------------------------------
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            buf = []
            i += 1
            while i < len(lines) and not lines[i].lstrip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            if lang.lower() in ("mermaid",):
                p = doc.add_paragraph()
                run = p.add_run(
                    "[Diagram — see Markdown source for the Mermaid definition. "
                    "Render in any Markdown viewer that supports Mermaid, "
                    "or paste the block at https://mermaid.live to export an image.]"
                )
                run.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            else:
                add_code_block(doc, "\n".join(buf), language=lang)
            continue

        # ---- table -------------------------------------------------------
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(
            r"^\|[\s\-\|:]+\|$", lines[i + 1].strip()
        ):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2  # skip header + separator
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(
                    [c.strip() for c in lines[i].strip().strip("|").split("|")]
                )
                i += 1
            tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
            tbl.style = "Light Grid Accent 1"
            for c, txt in enumerate(header):
                cell = tbl.rows[0].cells[c]
                cell.text = ""
                add_inline(cell.paragraphs[0], txt, base_bold=True)
                set_shading(cell, "DCE6F1")
            for r_i, row in enumerate(rows, start=1):
                for c, txt in enumerate(row):
                    cell = tbl.rows[r_i].cells[c]
                    cell.text = ""
                    add_inline(cell.paragraphs[0], txt)
            doc.add_paragraph()
            continue

        # ---- headings ----------------------------------------------------
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # strip anchor links like "Title {#anchor}"
            text = re.sub(r"\s*\{#.+?\}\s*$", "", text)
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        # ---- image -------------------------------------------------------
        m = re.match(r"^!\[(?P<alt>[^\]]*)\]\((?P<src>[^)]+)\)\s*$", stripped)
        if m:
            src = m.group("src")
            alt = m.group("alt")
            img_path = ROOT / src
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if img_path.exists():
                try:
                    p.add_run().add_picture(str(img_path), width=Inches(6.0))
                except Exception as e:
                    p.add_run(f"[Image: {src} — {e}]").italic = True
            else:
                run = p.add_run(f"[ Insert screenshot: {src} ]")
                run.italic = True
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            if alt:
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run(alt)
                run.italic = True
                run.font.size = Pt(9)
            i += 1
            continue

        # ---- block quote -------------------------------------------------
        if stripped.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, stripped[2:])
            i += 1
            continue

        # ---- bullet list -------------------------------------------------
        if re.match(r"^[\-\*]\s+", stripped):
            while i < len(lines) and re.match(r"^[\-\*]\s+", lines[i].strip()):
                txt = re.sub(r"^[\-\*]\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, txt)
                i += 1
            continue

        # ---- numbered list ----------------------------------------------
        if re.match(r"^\d+\.\s+", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                txt = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Number")
                add_inline(p, txt)
                i += 1
            continue

        # ---- blank line --------------------------------------------------
        if not stripped:
            i += 1
            continue

        # ---- regular paragraph (collect lines until blank / structural) -
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt.startswith("#")
                or nxt.startswith("```")
                or nxt.startswith("|")
                or nxt.startswith("> ")
                or nxt.startswith("![")
                or re.match(r"^[\-\*]\s+", nxt)
                or re.match(r"^\d+\.\s+", nxt)
                or nxt == "---"
                or nxt.startswith('<div')
                or nxt.startswith("</div>")
            ):
                break
            para_lines.append(nxt)
            i += 1
        para_text = " ".join(para_lines)
        p = doc.add_paragraph()
        add_inline(p, para_text)


def main() -> None:
    doc = Document()

    # Base styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    md_text = SRC.read_text(encoding="utf-8")
    parse_and_write(md_text, doc)

    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
